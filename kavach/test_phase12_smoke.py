"""test_phase12_smoke.py — End-to-end regression smoke test of ALL 7 major flows."""

import json
from pathlib import Path
import sys
import docx
from fastapi.testclient import TestClient

ROOT_DIR = Path(__file__).resolve().parent
if str(ROOT_DIR) not in sys.path:
    sys.path.insert(0, str(ROOT_DIR))

from backend import config
from backend.audit.logbook import read_events
from backend.guard.approve import get_approval
from backend.main import app

client = TestClient(app)

print("=== KAVACH Phase 12: 7-Flow End-to-End Regression Smoke Test ===")

# Flow 1: RAG Search with Citation
print("\n[Flow 1/7] Testing RAG Search with Citation...")
res_1 = client.post("/run", json={"task": "Search the SOPs for who must be notified during a Severity 1 incident."})
assert res_1.status_code == 200
data_1 = res_1.json()
assert data_1.get("status") == "complete"
assert len(data_1.get("sources", [])) > 0, "Expected at least 1 source citation"
print(f"  Result preview: {data_1.get('result', '')[:100]}...")
print(f"  Sources cited : {[s['filename'] for s in data_1.get('sources', [])]}")
print("  [PASS] Flow 1 Passed.")

# Flow 2: Docx Generation
print("\n[Flow 2/7] Testing Docx Generation...")
from backend.tools.writer import write_document
write_res = write_document(
    topic_or_content="Phase 12 System Verification and Health Note",
    sources=["incident_response_sop.md"],
    is_grounded=True,
    task_id="phase12-smoke-flow2",
)
doc_path = Path(write_res["file_path"])
assert doc_path.exists()
doc = docx.Document(str(doc_path))
assert len(doc.paragraphs) > 2
print(f"  Rendered docx: {doc_path.name} ({len(doc.paragraphs)} paragraphs)")
print("  [PASS] Flow 2 Passed.")

# Flow 3: Code Sandbox with Error Feedback Self-Correction
print("\n[Flow 3/7] Testing Code Sandbox with Error Feedback Self-Correction...")
from backend.tools.sandbox import _docker_available
if _docker_available():
    res_3 = client.post("/run", json={"task": "Write and run a python script that prints 'KAVACH_SANDBOX_ONLINE' and the value of 14 * 7"})
    assert res_3.status_code == 200
    data_3 = res_3.json()
    assert data_3.get("status") == "complete"
    assert any("98" in str(cr.get("stdout", "")) for cr in data_3.get("code_runs", [])) or "98" in str(data_3.get("result", ""))
    print(f"  Sandbox executed inside Docker: {len(data_3.get('code_runs', []))} code run(s)")
    print("  [PASS] Flow 3 Passed (Live container execution).")
else:
    res_3 = client.post("/run", json={"task": "Write and run a python script that prints 'KAVACH_SANDBOX_ONLINE' and the value of 14 * 7"})
    assert res_3.status_code == 200
    data_3 = res_3.json()
    assert any("Docker is not" in str(cr.get("stderr", "")) for cr in data_3.get("code_runs", []))
    print("  [NOTICE] Docker Desktop is not currently running on host.")
    print("  [PASS] Flow 3 Passed (Security Guard: host execution strictly prevented, fail-closed posture verified).")

# Flow 4: OCR Ingestion
print("\n[Flow 4/7] Testing OCR Ingestion...")
from backend.tools.ocr import extract_text
from PIL import Image, ImageDraw

img_path = config.OUTPUTS_DIR / "smoke_ocr_test.png"
img = Image.new("RGB", (320, 80), color=(255, 255, 255))
d = ImageDraw.Draw(img)
d.text((15, 30), "VALVE-402-ALPHA", fill=(0, 0, 0))
img.save(str(img_path))
ocr_res = extract_text(str(img_path), task_id="phase12-smoke-flow4")
print(f"  OCR extracted: '{ocr_res['text'].strip()}', engine={ocr_res['engine']}, confidence={ocr_res['confidence']:.2f}")
assert "VALVE" in ocr_res["text"].upper() or "402" in ocr_res["text"]
print("  [PASS] Flow 4 Passed.")

# Flow 5: Calculator with Steps
print("\n[Flow 5/7] Testing Calculator with Steps...")
from backend.tools.calc import calculate
calc_res = calculate(
    "Calculate the remaining pipe life where current thickness is 12.5 mm, minimum required thickness is 8.0 mm, and corrosion rate is 0.4 mm/year.",
    task_id="phase12-smoke-flow5",
)
assert calc_res.get("success") is True
assert calc_res.get("result") == 11.25
assert len(calc_res.get("steps", [])) >= 2
print(f"  Formula: {calc_res.get('formula_name')}")
print(f"  Result : {calc_res.get('result')} {calc_res.get('unit')}")
print(f"  Steps  : {calc_res.get('steps')}")
print("  [PASS] Flow 5 Passed.")

# Flow 6: Approval Gate (Approve Path)
print("\n[Flow 6/7] Testing Approval Gate (Approve Path)...")
task_id_6 = "phase12-smoke-flow6"
res_6 = client.post("/run", json={"task": "Draft an emergency containment procedure document for pipeline Bravo.", "task_id": task_id_6})
assert res_6.status_code == 200
data_6 = res_6.json()
assert data_6.get("status") == "awaiting_approval"
assert len(data_6.get("generated_files", [])) == 0, "No file should exist before approval"

# Submit Approve
appr_res = client.post(f"/approval/{task_id_6}", json={"decision": "approve"})
assert appr_res.status_code == 200
appr_data = appr_res.json()
assert appr_data.get("status") == "approved"
assert Path(appr_data["file_path"]).exists()
print(f"  Paused at risk: {data_6.get('approval', {}).get('risk')}")
print(f"  Approved and rendered: {appr_data.get('filename')}")
print("  [PASS] Flow 6 Passed.")

# Flow 7: Sovereignty Monitor & Zero External Calls
print("\n[Flow 7/7] Testing Sovereignty Audit Logging (Zero External Calls)...")
events = read_events()
external_calls_total = sum(e.get("external_calls", 0) for e in events)
print(f"  Total audit log events inspected: {len(events)}")
print(f"  Total external calls recorded   : {external_calls_total}")
assert external_calls_total == 0, f"Found {external_calls_total} non-zero external calls!"
print("  [PASS] Flow 7 Passed. 100% offline local sovereignty confirmed.")

print("\n=======================================================")
print("ALL 7 MAJOR FLOWS PASSED REGRESSION SMOKE TEST!")
print("=======================================================")
